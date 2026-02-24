"""
Document Ingestion Service
Handles PDF/DOCX upload, text extraction, chunking, embedding, and FAISS indexing.

Chunking Strategy:
  - Sliding window: 500-800 tokens with 100-token overlap
  - Preserves semantic context across chunk boundaries
"""

import os
import uuid
import logging
import numpy as np
from datetime import datetime
from typing import List, Tuple
from pathlib import Path

from sqlalchemy.orm import Session
from models.db_models import KnowledgeBaseDocument
from embeddings.embedding_service import embedding_service
from embeddings.vector_store import vector_store, DocumentChunk
from config import settings

logger = logging.getLogger(__name__)

UPLOAD_DIR = "./data/uploads"
os.makedirs(UPLOAD_DIR, exist_ok=True)


class IngestService:
    """
    End-to-end document ingestion pipeline:
      1. Save uploaded file to disk
      2. Extract text (PDF → PyMuPDF, DOCX → python-docx, TXT → direct)
      3. Chunk text with sliding window
      4. Generate embeddings (batch)
      5. Store in FAISS vector store
      6. Update KnowledgeBaseDocument record
    """

    def __init__(self):
        self.chunk_size = settings.CHUNK_SIZE        # chars (approximated)
        self.chunk_overlap = settings.CHUNK_OVERLAP

    # ──────────────────────────────────────────────────────────────────────────
    # Text Extraction
    # ──────────────────────────────────────────────────────────────────────────

    def _extract_text_pdf(self, filepath: str) -> str:
        """Extract text from PDF using PyMuPDF (fitz)."""
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = "\n".join(page.get_text() for page in doc)
            doc.close()
            return text
        except ImportError:
            logger.error("PyMuPDF not installed. Run: pip install pymupdf")
            raise

    def _extract_text_docx(self, filepath: str) -> str:
        """Extract text from DOCX using python-docx."""
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join(para.text for para in doc.paragraphs if para.text.strip())
        except ImportError:
            logger.error("python-docx not installed. Run: pip install python-docx")
            raise

    def _extract_text(self, filepath: str, file_type: str) -> str:
        """Dispatch to appropriate extractor based on file type."""
        if file_type == "pdf":
            return self._extract_text_pdf(filepath)
        elif file_type == "docx":
            return self._extract_text_docx(filepath)
        elif file_type == "txt":
            with open(filepath, "r", encoding="utf-8") as f:
                return f.read()
        else:
            raise ValueError(f"Unsupported file type: {file_type}")

    # ──────────────────────────────────────────────────────────────────────────
    # Chunking
    # ──────────────────────────────────────────────────────────────────────────

    def _chunk_text(self, text: str) -> List[str]:
        """
        Sliding window chunking with overlap.
        Character-based approximation: 1 token ≈ 4 characters.
        chunk_size=600 tokens → ~2400 chars
        """
        char_chunk = self.chunk_size * 4
        char_overlap = self.chunk_overlap * 4

        chunks = []
        start = 0

        while start < len(text):
            end = start + char_chunk
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            start += char_chunk - char_overlap

        logger.info(f"Text chunked into {len(chunks)} pieces (chunk_size={self.chunk_size} tokens)")
        return chunks

    # ──────────────────────────────────────────────────────────────────────────
    # Main Ingest Pipeline
    # ──────────────────────────────────────────────────────────────────────────

    async def ingest_document(
        self,
        file_bytes: bytes,
        original_filename: str,
        db: Session,
        uploaded_by: str = "system",
    ) -> KnowledgeBaseDocument:
        """
        Full ingestion pipeline for an uploaded document.

        Args:
            file_bytes: Raw file bytes from upload
            original_filename: Original name of the uploaded file
            db: SQLAlchemy session
            uploaded_by: Username or email of uploader

        Returns:
            Updated KnowledgeBaseDocument record
        """
        ext = Path(original_filename).suffix.lstrip(".").lower()
        if ext not in ("pdf", "docx", "txt"):
            raise ValueError(f"Unsupported file type: .{ext}. Supported: pdf, docx, txt")

        # ── Save file ─────────────────────────────────────────────────────────
        safe_name = f"{uuid.uuid4().hex}_{original_filename}"
        filepath = os.path.join(UPLOAD_DIR, safe_name)
        with open(filepath, "wb") as f:
            f.write(file_bytes)

        # ── Create DB record ──────────────────────────────────────────────────
        doc_record = KnowledgeBaseDocument(
            filename=safe_name,
            original_filename=original_filename,
            file_type=ext,
            status="processing",
            file_size_bytes=len(file_bytes),
            uploaded_by=uploaded_by,
        )
        db.add(doc_record)
        db.commit()
        db.refresh(doc_record)

        try:
            # ── Extract text ──────────────────────────────────────────────────
            text = self._extract_text(filepath, ext)
            if not text.strip():
                raise ValueError("Extracted text is empty. Document may be image-only or encrypted.")

            # ── Chunk ─────────────────────────────────────────────────────────
            chunks_text = self._chunk_text(text)

            # ── Embed (batch) ─────────────────────────────────────────────────
            embeddings = embedding_service.encode(chunks_text).astype(np.float32)

            # ── Build DocumentChunk objects ───────────────────────────────────
            doc_chunks = [
                DocumentChunk(
                    chunk_id=str(uuid.uuid4()),
                    content=chunk,
                    source=original_filename,
                    chunk_index=i,
                    doc_id=doc_record.id,
                )
                for i, chunk in enumerate(chunks_text)
            ]

            # ── Store in FAISS ────────────────────────────────────────────────
            vector_store.add_chunks(embeddings, doc_chunks)

            # ── Update DB record ──────────────────────────────────────────────
            doc_record.chunk_count = len(doc_chunks)
            doc_record.status = "indexed"
            doc_record.indexed_at = datetime.utcnow()

        except Exception as e:
            logger.error(f"Ingestion failed for {original_filename}: {e}")
            doc_record.status = "failed"
            raise
        finally:
            db.commit()
            db.refresh(doc_record)

        logger.info(
            f"Successfully ingested '{original_filename}': "
            f"{doc_record.chunk_count} chunks indexed (doc_id={doc_record.id})"
        )
        return doc_record


# Singleton
ingest_service = IngestService()
